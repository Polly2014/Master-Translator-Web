#!/usr/bin/env python3
"""
Master Translator Web Application
智能分块翻译的 Web 界面 - Hackathon Demo 版
"""

import os
import json
import time
import re
from pathlib import Path
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file
from flask_socketio import SocketIO, emit
from flask_cors import CORS
from werkzeug.utils import secure_filename
import threading
from litellm import completion
import os
from dotenv import load_dotenv
from docx import Document
from markdownify import markdownify as md

# 自动加载 .env 文件（如果存在）
load_dotenv()

# ============ Flask 配置 ============
app = Flask(__name__)
app.config['SECRET_KEY'] = 'master-translator-secret-2024'
app.config['UPLOAD_FOLDER'] = Path('./uploads')
app.config['OUTPUT_FOLDER'] = Path('./outputs')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB

CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

# 确保目录存在
app.config['UPLOAD_FOLDER'].mkdir(exist_ok=True)
app.config['OUTPUT_FOLDER'].mkdir(exist_ok=True)

# ============ 翻译配置 ============
# 从环境变量读取 OpenRouter API Key，避免硬编码泄漏
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
if not OPENROUTER_API_KEY:
    print("[WARN] OPENROUTER_API_KEY 未设置，翻译接口将不可用。请在本地创建 .env 或通过 shell 导出变量。")

# ============ 模型配置字典 ============
MODEL_CONFIGS = {
    'deepseek-free': {
        'name': 'tngtech/deepseek-r1t-chimera:free',
        'max_tokens': 16000,
        'temperature': 0.3,
        'cost_per_1k': 0.0,
        'description': '免费模型，适合 Demo 和开发测试',
        'speed': 'fast',
        'quality': 'good'
    },
    'claude-sonnet-4': {
        'name': 'anthropic/claude-sonnet-4',
        'max_tokens': 100000,
        'temperature': 0.3,
        'cost_per_1k': 0.01,
        'description': '最高质量，适合生产环境',
        'speed': 'medium',
        'quality': 'excellent'
    },
    'gpt-4o': {
        'name': 'openai/gpt-4o',
        'max_tokens': 100000,
        'temperature': 0.3,
        'cost_per_1k': 0.0067,
        'description': '平衡性能和成本',
        'speed': 'fast',
        'quality': 'excellent'
    },
    'deepseek-v3': {
        'name': 'deepseek/deepseek-chat',
        'max_tokens': 64000,
        'temperature': 0.3,
        'cost_per_1k': 0.0013,
        'description': '高性价比，适合大规模生产',
        'speed': 'very-fast',
        'quality': 'good'
    }
}

# 当前使用的模型（修改这里切换模型)
ACTIVE_MODEL = 'deepseek-free'  # 可选: deepseek-free, claude-sonnet-4, gpt-4o, deepseek-v3

# 从配置中加载当前模型参数
current_config = MODEL_CONFIGS[ACTIVE_MODEL]
MODEL = current_config['name']
MAX_TOKENS = current_config['max_tokens']
TEMPERATURE = current_config['temperature']
TIMEOUT = 3600

# ============ 分块配置 ============
# 使用专门的 Demo 文件 (demo_files/) 进行演示
# Ultra Quick Demo: 200 words (~20-30s, 3 chunks)
# Standard Demo: 1,037 words (~3-5min)
DEMO_MODE = True  # 已有专门 Demo 文件，使用生产配置

if DEMO_MODE:
    CHUNK_TARGET_SIZE = 800          # Demo: 超小块，确保 Quick Demo 3章→3块（~20-30s）
    CONTEXT_PARAGRAPHS = 1           # Demo: 减少上下文，加快速度
    OVERLAP_CHECK_CHARS = 100        # Demo: 减少重叠检查
else:
    CHUNK_TARGET_SIZE = 110000       # 生产: 大块，减少 API 调用
    CONTEXT_PARAGRAPHS = 2           # 生产: 更多上下文，提高质量
    OVERLAP_CHECK_CHARS = 200        # 生产: 更多重叠检查

# 其他配置

LANGUAGES = {
    # East Asian
    'Chinese': '中文 (简体)',
    'Traditional Chinese': '中文 (繁体)',
    'Japanese': '日语',
    'Korean': '韩语',
    
    # European
    'French': '法语',
    'German': '德语',
    'Spanish': '西班牙语',
    'Italian': '意大利语',
    'Portuguese': '葡萄牙语',
    'Russian': '俄语',
    'Polish': '波兰语',
    'Dutch': '荷兰语',
    'Swedish': '瑞典语',
    
    # Middle Eastern & South Asian
    'Arabic': '阿拉伯语',
    'Hebrew': '希伯来语',
    'Hindi': '印地语',
    'Urdu': '乌尔都语',
    'Persian': '波斯语',
    'Turkish': '土耳其语',
    
    # Southeast Asian
    'Thai': '泰语',
    'Vietnamese': '越南语',
    'Indonesian': '印尼语',
    'Malay': '马来语',
    'Spanish': '西班牙语',
    'German': '德语'
}

# ============ 全局任务管理 ============
tasks = {}  # 存储所有翻译任务


class TranslationTask:
    """翻译任务类"""
    def __init__(self, task_id, filename, language):
        self.task_id = task_id
        self.filename = filename
        self.language = language
        self.status = 'pending'  # pending, analyzing, translating, completed, failed
        self.progress = 0
        self.current_chunk = 0
        self.total_chunks = 0
        self.chunks_info = []
        self.logs = []
        self.source_content = ""
        self.result_file = None
        self.start_time = None
        self.end_time = None
        self.error = None
        self.use_terminology = True  # 默认使用术语数据库
        
    def emit_log(self, message, level='info', update_last=False):
        """发送日志到前端
        
        Args:
            message: 日志消息
            level: 日志级别 (info, success, error, warning, progress)
            update_last: 是否更新最后一条日志（用于进度更新)
        """
        log_entry = {
            'timestamp': datetime.now().strftime('%H:%M:%S'),
            'message': message,
            'level': level,
            'update_last': update_last
        }
        self.logs.append(log_entry)
        socketio.emit('log', log_entry, room=self.task_id)
        
    def emit_progress(self, progress, chunk_progress=0):
        """发送进度到前端"""
        self.progress = progress
        socketio.emit('progress', {
            'overall': progress,
            'chunk': chunk_progress,
            'current_chunk': self.current_chunk,
            'total_chunks': self.total_chunks
        }, room=self.task_id)


# ============ 翻译核心函数（改造自 script_v3_chunked.py)============

def convert_docx_to_markdown(docx_path):
    """
    将 Word 文档转换为 Markdown 格式
    
    Args:
        docx_path: Word 文档路径
        
    Returns:
        str: Markdown 格式的文本内容
    """
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        for para in doc.paragraphs:
            # 处理标题
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                markdown_content.append(f"{'#' * level} {para.text}\n")
            # 处理普通段落
            elif para.text.strip():
                markdown_content.append(f"{para.text}\n")
        
        # 处理表格
        for table in doc.tables:
            markdown_content.append("\n")
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                markdown_content.append("| " + " | ".join(cells) + " |")
                # 添加表头分隔符
                if i == 0:
                    markdown_content.append("| " + " | ".join(["---"] * len(cells)) + " |")
            markdown_content.append("\n")
        
        result = "\n".join(markdown_content)
        
        # 清理多余的空行
        result = re.sub(r'\n{3,}', '\n\n', result)
        
        return result.strip()
        
    except Exception as e:
        raise Exception(f"DOCX conversion failed: {str(e)}")


def extract_chapters(content):
    """提取章节结构"""
    pattern = r'^(#{1,2}) (.+)$'
    chapters = []
    lines = content.split('\n')
    
    for i, line in enumerate(lines):
        match = re.match(pattern, line)
        if match:
            level = len(match.group(1))
            title = match.group(2).strip()
            start_pos = sum(len(l) + 1 for l in lines[:i])
            chapters.append({
                'line': i + 1,
                'level': level,
                'title': title,
                'start_pos': start_pos
            })
    
    total_len = len(content)
    for i, ch in enumerate(chapters):
        if i < len(chapters) - 1:
            ch['chars'] = chapters[i+1]['start_pos'] - ch['start_pos']
            ch['end_pos'] = chapters[i+1]['start_pos']
        else:
            ch['chars'] = total_len - ch['start_pos']
            ch['end_pos'] = total_len
    
    return chapters


def plan_chunks(chapters, content):
    """规划分Chunk - 支持任何 Markdown 文件结构"""
    chunks = []
    current_chunk_chapters = []
    current_size = 0
    
    # 优先使用 level 2 标题，如果没有则使用 level 1
    main_chapters = [c for c in chapters if c['level'] == 2]
    if not main_chapters:
        main_chapters = [c for c in chapters if c['level'] == 1]
    
    # 如果仍然没有章节，按固定大小分块整个文档
    if not main_chapters:
        total_len = len(content)
        num_chunks = max(1, (total_len + CHUNK_TARGET_SIZE - 1) // CHUNK_TARGET_SIZE)
        for i in range(num_chunks):
            start = i * CHUNK_TARGET_SIZE
            end = min(start + CHUNK_TARGET_SIZE, total_len)
            chunks.append({
                'id': i + 1,
                'chapters': [f'Segment {i+1}'],
                'start_pos': start,
                'end_pos': end,
                'size': end - start,
                'content': content[start:end]
            })
        return chunks
    
    if main_chapters:
        first_chapter_pos = main_chapters[0]['start_pos']
        prologue_content = content[:first_chapter_pos].strip()
        prologue_size = len(prologue_content)
    else:
        prologue_content = ""
        prologue_size = 0
    
    if main_chapters:
        last_chapter_pos = main_chapters[-1]['end_pos']
        epilogue_content = content[last_chapter_pos:].strip()
        epilogue_size = len(epilogue_content)
    else:
        epilogue_content = ""
        epilogue_size = 0
    
    for chapter in main_chapters:
        if current_size > 0 and current_size + chapter['chars'] > CHUNK_TARGET_SIZE:
            chunk_start = current_chunk_chapters[0]['start_pos']
            chunk_end = current_chunk_chapters[-1]['end_pos']
            chunks.append({
                'id': len(chunks) + 1,
                'chapters': [c['title'] for c in current_chunk_chapters],
                'start_pos': chunk_start,
                'end_pos': chunk_end,
                'size': current_size,
                'content': content[chunk_start:chunk_end]
            })
            current_chunk_chapters = []
            current_size = 0
        
        current_chunk_chapters.append(chapter)
        current_size += chapter['chars']
    
    if current_chunk_chapters:
        chunk_start = current_chunk_chapters[0]['start_pos']
        chunk_end = current_chunk_chapters[-1]['end_pos']
        chunks.append({
            'id': len(chunks) + 1,
            'chapters': [c['title'] for c in current_chunk_chapters],
            'start_pos': chunk_start,
            'end_pos': chunk_end,
            'size': current_size,
            'content': content[chunk_start:chunk_end]
        })
    
    if chunks and prologue_content:
        chunks[0]['content'] = prologue_content + "\n\n" + chunks[0]['content']
        chunks[0]['size'] += prologue_size + 2
        chunks[0]['has_prologue'] = True
        chunks[0]['prologue_size'] = prologue_size
    
    if chunks and epilogue_content:
        chunks[-1]['content'] = chunks[-1]['content'] + "\n\n" + epilogue_content
        chunks[-1]['size'] += epilogue_size + 2
        chunks[-1]['has_epilogue'] = True
        chunks[-1]['epilogue_size'] = epilogue_size
    
    return chunks


# ============ 模型管理函数 ============

def get_model_info():
    """获取当前模型配置信息"""
    config = MODEL_CONFIGS[ACTIVE_MODEL]
    return {
        'active_model': ACTIVE_MODEL,
        'model_name': config['name'],
        'max_tokens': config['max_tokens'],
        'temperature': config['temperature'],
        'cost_per_1k': config['cost_per_1k'],
        'description': config['description'],
        'speed': config['speed'],
        'quality': config['quality']
    }


def list_available_models():
    """列出所有可用模型"""
    return {
        key: {
            'name': config['name'],
            'description': config['description'],
            'cost': f"${config['cost_per_1k']:.4f}/1K chars",
            'speed': config['speed'],
            'quality': config['quality']
        }
        for key, config in MODEL_CONFIGS.items()
    }


def load_terminology_db():
    """加载精选术语数据库"""
    term_file = Path(__file__).parent / 'terminology_curated.json'
    
    if term_file.exists():
        with open(term_file, 'r', encoding='utf-8') as f:
            terminology = json.load(f)
        
        all_terms = []
        for category, terms in terminology.items():
            all_terms.extend(terms)
        
        return all_terms
    return None


def extract_terminology_from_chunk(translation_text, source_text):
    """从翻译块中提取新术语（动态提取)"""
    import re
    
    # 1. 提取专有名词（大写开头的词组)
    proper_nouns = re.findall(r'\b[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){0,3}\b', source_text)
    
    # 过滤常见词
    common_words = {
        'The', 'This', 'That', 'These', 'Those', 'Chapter', 'Part', 'Section',
        'And', 'But', 'For', 'With', 'When', 'Where', 'Which', 'What', 'How',
        'Figure', 'Table', 'Some', 'Many', 'Most', 'All', 'Each', 'Every',
        'First', 'Second', 'Third', 'Last', 'Next', 'Previous', 'Introduction',
        'Conclusion', 'Summary', 'Overview', 'Today', 'Tomorrow', 'Yesterday'
    }
    proper_nouns = [t for t in set(proper_nouns) if t not in common_words and len(t) > 2]
    
    # 2. 检测技术术语（常见技术词汇在源文本中出现)
    tech_keywords = [
        'AI', 'ML', 'API', 'GPU', 'CPU', 'DNA', 'RNA', 'AGI',
        'artificial intelligence', 'machine learning', 'deep learning',
        'neural network', 'algorithm', 'model', 'dataset', 'training',
        'inference', 'transformer', 'attention', 'backpropagation',
        'reinforcement learning', 'supervised learning', 'unsupervised learning',
        'natural language processing', 'computer vision', 'robotics',
        'blockchain', 'cryptocurrency', 'quantum computing',
        'biotechnology', 'synthetic biology', 'gene editing', 'CRISPR'
    ]
    
    found_tech = [term for term in tech_keywords if term.lower() in source_text.lower()]
    
    # 合并并去重
    extracted_terms = list(set(proper_nouns + found_tech))
    
    return extracted_terms


def get_context_from_previous(prev_translation):
    """从前一块翻译末尾提取上下文"""
    if not prev_translation:
        return ""
    
    paragraphs = [p.strip() for p in prev_translation.split('\n\n') if p.strip()]
    context_paras = paragraphs[-CONTEXT_PARAGRAPHS:] if len(paragraphs) >= CONTEXT_PARAGRAPHS else paragraphs
    context = '\n\n'.join(context_paras)
    
    return context


def translate_chunk_web(task, chunk_id, total_chunks, chunk_content, language, 
                        prev_context="", terminology=None):
    """Web版翻译单块（带实时日志)"""
    
    task.emit_log(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━", 'info')
    task.emit_log(f"🔄 Starting chunk {chunk_id}/{total_chunks}", 'info')
    task.emit_log(f"📝 Input size: {len(chunk_content):,} characters", 'info')
    
    # Demo 优化：快速显示开始信息
    if DEMO_MODE:
        task.emit_log(f"⚡ Demo mode: Using small chunks for quick demonstration", 'info')
    
    # 显示使用的术语数量
    if terminology:
        task.emit_log(f"� Using {len(terminology)} terms for consistency", 'info')
    
    system_prompt = f"""You are a professional book translator. Translate the following book excerpt from English to {language}.

CRITICAL REQUIREMENTS:
1. **Preserve ALL markdown formatting**: headers (#, ##, ###), lists, quotes (>), code blocks, links
2. **Maintain structure**: Keep all chapters, sections, paragraphs exactly as structured
3. **Preserve names**: Keep all person names, company names, book titles in original English
4. **Technical terms**: Translate technical terms accurately, add English in parentheses for first occurrence if needed
5. **Natural language**: Use native {language} expression, not word-by-word translation
6. **Completeness**: Translate EVERY sentence, don't skip any content
7. **Consistency**: Maintain consistent terminology throughout the book

CONTEXT: This is chunk {chunk_id}/{total_chunks} of the complete book.
"""
    
    context_info = ""
    if prev_context:
        context_info = f"""
<previous_context>
For continuity, here are the last paragraphs from the previous chunk:
{prev_context}
</previous_context>
"""
    
    term_info = ""
    if terminology:
        term_info = f"""
<key_terminology>
Important terms to maintain consistency:
{', '.join(terminology[:25])}
</key_terminology>
"""
    
    user_prompt = f"""{context_info}{term_info}

Now translate this section to {language}:

---BEGIN CONTENT---

{chunk_content}

---END CONTENT---"""

    try:
        start_time = time.time()
        translated_text = ""
        last_update = start_time
        
        response = completion(
            model=f"openrouter/{MODEL}",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            api_key=OPENROUTER_API_KEY,
            max_tokens=MAX_TOKENS,
            temperature=TEMPERATURE,
            timeout=TIMEOUT,
            stream=True,
            extra_headers={
                "HTTP-Referer": "https://github.com/Polly2014",
                "X-Title": "Master Translator Web"
            }
        )
        
        for chunk in response:
            if hasattr(chunk, 'choices') and len(chunk.choices) > 0:
                delta = chunk.choices[0].delta
                if hasattr(delta, 'content') and delta.content:
                    translated_text += delta.content
                    
                    # Demo 模式：更频繁更新（每1seconds)，生产模式：每5seconds
                    update_interval = 1 if DEMO_MODE else 5
                    chars_threshold = 2000 if DEMO_MODE else 10000
                    
                    now = time.time()
                    if now - last_update > update_interval or len(translated_text) % chars_threshold < 100:
                        elapsed = now - start_time
                        speed = len(translated_text) / elapsed if elapsed > 0 else 0
                        chunk_progress = min(95, int((len(translated_text) / (len(chunk_content) * 1.5)) * 100))
                        
                        task.emit_progress(
                            progress=int(((chunk_id - 1) + chunk_progress / 100) / total_chunks * 100),
                            chunk_progress=chunk_progress
                        )
                        # 使用 update_last=True 更新进度消息而不是追加
                        task.emit_log(f"📥 Receiving translation... {len(translated_text):,} characters ({speed:.0f} c/s)", 'progress', update_last=True)
                        last_update = now
        
        elapsed = time.time() - start_time
        speed = len(translated_text) / elapsed if elapsed > 0 else 0
        
        task.emit_log(f"✅ Chunk {chunk_id} completed: {len(translated_text):,} characters ({speed:.0f} c/s, {elapsed:.0f}s)", 'success')
        
        return translated_text
        
    except Exception as e:
        task.emit_log(f"❌ Chunk {chunk_id} translation failed: {str(e)}", 'error')
        raise


def translate_book_task(task):
    """执行翻译任务（后台线程)"""
    try:
        task.status = 'translating'
        task.start_time = datetime.now()
        task.emit_log(f"🚀 Starting translation task", 'info')
        task.emit_log(f"📚 File: {task.filename}", 'info')
        task.emit_log(f"🌍 Target language: {LANGUAGES.get(task.language, task.language)}", 'info')
        
        # 加载术语库（根据用户选择)
        terminology = None
        curated_count = 0
        if task.use_terminology:
            terminology = load_terminology_db()
            if terminology:
                curated_count = len(terminology)
                task.emit_log(f"📚 Loaded curated terminology: {curated_count} terms", 'success')
                task.emit_log(f"🔄 Hybrid mode: Will extract new terms dynamically after first chunk", 'info')
            else:
                task.emit_log(f"⚠️  Terminology database not found, will use pure dynamic extraction mode", 'warning')
                terminology = []
        else:
            task.emit_log(f"ℹ️  User chose not to use terminology database", 'info')
        
        # 翻译所有块
        all_translations = []
        prev_context = ""
        
        for chunk in task.chunks_info:
            task.current_chunk = chunk['id']
            
            translation = translate_chunk_web(
                task=task,
                chunk_id=chunk['id'],
                total_chunks=task.total_chunks,
                chunk_content=chunk['content'],
                language=task.language,
                prev_context=prev_context,
                terminology=terminology
            )
            
            all_translations.append({
                'chunk_id': chunk['id'],
                'translation': translation,
                'chapters': chunk['chapters']
            })
            
            # 🔥 关键：从第一块提取新术语（混合模式)
            if chunk['id'] == 1 and terminology is not None:
                task.emit_log(f"🔍 Extracting new terms from first chunk...", 'info')
                extracted_terms = extract_terminology_from_chunk(translation, chunk['content'])
                
                # 过滤已存在的术语
                new_terms = [t for t in extracted_terms if t not in terminology]
                
                if new_terms:
                    terminology.extend(new_terms)
                    task.emit_log(f"✨ Extracted {len(new_terms)} new terms (total: {len(terminology)})", 'success')
                    if len(new_terms) <= 10:
                        task.emit_log(f"   New: {', '.join(new_terms[:10])}", 'info')
                    else:
                        task.emit_log(f"   Sample: {', '.join(new_terms[:5])}...", 'info')
                else:
                    task.emit_log(f"✅ First chunk terms already covered, no supplement needed", 'success')
            
            # 更新上下文
            prev_context = get_context_from_previous(translation)
            
            # 增量保存
            output_file = app.config['OUTPUT_FOLDER'] / f"{task.task_id}_{task.language}.md"
            with open(output_file, 'w', encoding='utf-8') as f:
                for t in all_translations:
                    f.write(f"\n\n<!-- Chunk {t['chunk_id']}: {', '.join(t['chapters'])} -->\n\n")
                    f.write(t['translation'])
            
            task.emit_log(f"💾 Progress saved ({chunk['id']}/{task.total_chunks})", 'info')
        
        # 最终合并
        final_translation = '\n\n'.join(t['translation'] for t in all_translations)
        
        # 保存最终结果
        output_file = app.config['OUTPUT_FOLDER'] / f"{task.task_id}_{task.language}_final.md"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(final_translation)
        
        task.result_file = str(output_file)
        task.status = 'completed'
        task.end_time = datetime.now()
        task.progress = 100
        
        elapsed = (task.end_time - task.start_time).total_seconds()
        task.emit_log(f"🎉 Translation completed!", 'success')
        task.emit_log(f"📊 Total: {len(final_translation):,} characters", 'success')
        task.emit_log(f"⏱️  Time elapsed: {elapsed:.0f} seconds", 'success')
        task.emit_progress(100, 100)
        
    except Exception as e:
        task.status = 'failed'
        task.error = str(e)
        task.emit_log(f"💥 translation failed: {str(e)}", 'error')


# ============ Flask 路由 ============

@app.route('/')
def index():
    """主页"""
    return render_template('index.html')


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文件（支持 .md 和 .docx）"""
    if 'file' not in request.files:
        return jsonify({'error': '没有文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '文件名为空'}), 400
    
    # 检查文件类型
    file_ext = file.filename.rsplit('.', 1)[-1].lower()
    if file_ext not in ['md', 'docx']:
        return jsonify({'error': '只支持 Markdown (.md) 和 Word (.docx) 文件'}), 400
    
    # 保存文件
    filename = secure_filename(file.filename)
    task_id = f"{int(time.time())}_{filename.rsplit('.', 1)[0]}"
    filepath = app.config['UPLOAD_FOLDER'] / f"{task_id}_{filename}"
    file.save(filepath)
    
    # 读取内容
    try:
        if file_ext == 'docx':
            # 转换 DOCX 为 Markdown
            content = convert_docx_to_markdown(filepath)
            # 保存转换后的 Markdown 版本
            md_filepath = app.config['UPLOAD_FOLDER'] / f"{task_id}_converted.md"
            with open(md_filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            conversion_note = "✅ Word document converted to Markdown automatically"
        else:
            # 直接读取 Markdown
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            conversion_note = None
    except Exception as e:
        return jsonify({'error': f'File processing failed: {str(e)}'}), 500
    
    response_data = {
        'task_id': task_id,
        'filename': filename,
        'size': len(content),
        'chars': len(content),
        'words': len(content.split()),
        'file_type': file_ext
    }
    
    if conversion_note:
        response_data['conversion_note'] = conversion_note
    
    return jsonify(response_data)


@app.route('/api/analyze/<task_id>', methods=['POST'])
def analyze_file(task_id):
    """分析文件并规划分块"""
    data = request.json
    language = data.get('language', 'Japanese')
    
    # 查找上传的文件（优先查找转换后的 .md 文件）
    converted_file = app.config['UPLOAD_FOLDER'] / f"{task_id}_converted.md"
    if converted_file.exists():
        filepath = converted_file
    else:
        files = list(app.config['UPLOAD_FOLDER'].glob(f"{task_id}_*"))
        if not files:
            return jsonify({'error': '文件不存在'}), 404
        filepath = files[0]
    
    # 读取内容
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 分析章节和分块
    chapters = extract_chapters(content)
    chunks = plan_chunks(chapters, content)
    
    # 创建任务
    task = TranslationTask(task_id, filepath.name, language)
    task.source_content = content
    task.total_chunks = len(chunks)
    task.chunks_info = chunks
    task.status = 'analyzed'
    
    tasks[task_id] = task
    
    # 返回分块信息
    chunks_summary = []
    for chunk in chunks:
        chunks_summary.append({
            'id': chunk['id'],
            'size': chunk['size'],
            'chapters': chunk['chapters'],
            'has_prologue': chunk.get('has_prologue', False),
            'has_epilogue': chunk.get('has_epilogue', False)
        })
    
    return jsonify({
        'task_id': task_id,
        'total_chunks': len(chunks),
        'total_chars': len(content),
        'chunks': chunks_summary
    })


@app.route('/api/translate/<task_id>', methods=['POST'])
def start_translation(task_id):
    """开始翻译"""
    if task_id not in tasks:
        return jsonify({'error': '任务不存在'}), 404
    
    task = tasks[task_id]
    
    if task.status != 'analyzed':
        return jsonify({'error': '任务状态错误'}), 400
    
    # 获取配置选项
    data = request.json or {}
    task.use_terminology = data.get('use_terminology', True)
    
    # 在后台线程中执行翻译
    thread = threading.Thread(target=translate_book_task, args=(task,))
    thread.daemon = True
    thread.start()
    
    return jsonify({'status': 'started', 'task_id': task_id})


@app.route('/api/status/<task_id>')
def get_status(task_id):
    """获取任务状态"""
    if task_id not in tasks:
        return jsonify({'error': '任务不存在'}), 404
    
    task = tasks[task_id]
    
    return jsonify({
        'task_id': task_id,
        'status': task.status,
        'progress': task.progress,
        'current_chunk': task.current_chunk,
        'total_chunks': task.total_chunks,
        'result_file': task.result_file,
        'error': task.error
    })


@app.route('/api/download/<task_id>')
def download_result(task_id):
    """下载翻译结果"""
    if task_id not in tasks:
        return jsonify({'error': '任务不存在'}), 404
    
    task = tasks[task_id]
    
    if task.status != 'completed' or not task.result_file:
        return jsonify({'error': '翻译未完成'}), 400
    
    return send_file(
        task.result_file,
        as_attachment=True,
        download_name=f"{task.filename.split('.')[0]}_{task.language}.md"
    )


@app.route('/api/preview/<task_id>')
def preview_result(task_id):
    """预览翻译结果内容"""
    if task_id not in tasks:
        return jsonify({'error': '任务不存在'}), 404
    
    task = tasks[task_id]
    
    if task.status != 'completed' or not task.result_file:
        return jsonify({'error': '翻译未完成'}), 400
    
    try:
        with open(task.result_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return jsonify({
            'success': True,
            'content': content,
            'filename': f"{task.filename.split('.')[0]}_{task.language}.md",
            'language': task.language
        })
    except Exception as e:
        return jsonify({'error': f'读取文件失败: {str(e)}'}), 500


@app.route('/api/preview-source/<task_id>')
def preview_source(task_id):
    """预览上传的源文件内容"""
    if task_id not in tasks:
        return jsonify({'error': '任务不存在'}), 404
    
    task = tasks[task_id]
    
    if not task.source_content:
        return jsonify({'error': '源文件内容不可用'}), 400
    
    try:
        return jsonify({
            'success': True,
            'content': task.source_content,
            'filename': task.filename,
            'language': 'English (source)'
        })
    except Exception as e:
        return jsonify({'error': f'读取源文件失败: {str(e)}'}), 500


@app.route('/api/preview-chunk/<task_id>/<int:chunk_id>')
def preview_chunk(task_id, chunk_id):
    """预览特定 chunk 的内容"""
    if task_id not in tasks:
        return jsonify({'error': '任务不存在'}), 404
    
    task = tasks[task_id]
    
    if not task.chunks_info:
        return jsonify({'error': 'Chunk 信息不可用，请先分析文件'}), 400
    
    # 查找对应的 chunk
    chunk_data = None
    for chunk in task.chunks_info:
        if chunk['id'] == chunk_id:
            chunk_data = chunk
            break
    
    if not chunk_data:
        return jsonify({'error': f'Chunk {chunk_id} 不存在'}), 404
    
    try:
        # 获取 chunk 的完整内容
        content = chunk_data.get('content', '')
        
        if not content:
            # 如果没有保存内容，尝试从源文件中提取
            if chunk_data.get('start_pos') is not None and chunk_data.get('end_pos') is not None:
                content = task.source_content[chunk_data['start_pos']:chunk_data['end_pos']]
            else:
                content = f"# Chunk {chunk_id}\n\nContent not available"
        
        return jsonify({
            'success': True,
            'content': content,
            'chunk_id': chunk_id,
            'chapters': chunk_data.get('chapters', []),
            'size': chunk_data.get('size', 0),
            'has_prologue': chunk_data.get('has_prologue', False),
            'has_epilogue': chunk_data.get('has_epilogue', False)
        })
    except Exception as e:
        return jsonify({'error': f'读取 Chunk 失败: {str(e)}'}), 500


@app.route('/api/terminology')
def get_terminology():
    """获取术语数据库（包含动态提取说明)"""
    term_file = Path(__file__).parent / 'terminology_curated.json'
    
    if not term_file.exists():
        return jsonify({
            'error': 'Terminology database not found',
            'path': str(term_file),
            'note': 'Will use dynamic extraction from first chunk if enabled'
        }), 404
    
    try:
        with open(term_file, 'r', encoding='utf-8') as f:
            terminology = json.load(f)
        
        # 统计信息
        stats = {
            'total': sum(len(terms) for terms in terminology.values()),
            'categories': {}
        }
        
        for category, terms in terminology.items():
            stats['categories'][category] = len(terms)
        
        return jsonify({
            'terminology': terminology,
            'stats': stats,
            'mode': 'hybrid',
            'description': 'Curated terms + dynamic extraction from first chunk'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/model-info')
def get_model_info_api():
    """获取当前模型配置信息"""
    try:
        info = get_model_info()
        return jsonify({
            'success': True,
            'model_info': info
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/models')
def get_models_list():
    """获取所有可用模型列表"""
    try:
        models = list_available_models()
        return jsonify({
            'success': True,
            'active_model': ACTIVE_MODEL,
            'models': models
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ============ WebSocket 事件 ============

@socketio.on('connect')
def handle_connect():
    """客户端连接"""
    print(f"✅ Client connected")
    return True


@socketio.on('disconnect')
def handle_disconnect():
    """客户端断开"""
    print(f"❌ Client disconnected")
    return True


@socketio.on('join')
def handle_join(data):
    """加入任务房间"""
    task_id = data.get('task_id')
    if task_id:
        # 使用 Flask-SocketIO 的 join_room
        from flask_socketio import join_room
        join_room(task_id)
        print(f"📥 Client joined room: {task_id}")
        return {'status': 'joined', 'task_id': task_id}
    return {'status': 'error', 'message': 'No task_id provided'}


# ============ 启动服务 ============

if __name__ == '__main__':
    print(f"""
{'='*80}
🚀 Master Translator Web - Hackathon Demo
{'='*80}
📡 服务器启动在: http://localhost:5001
🔌 WebSocket 已启用
📁 上传目录: {app.config['UPLOAD_FOLDER']}
📁 输出目录: {app.config['OUTPUT_FOLDER']}
{'='*80}
""")
    socketio.run(app, debug=True, host='0.0.0.0', port=5001)
